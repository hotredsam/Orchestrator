#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import json
import math
import os
import random
import shutil
import statistics
import subprocess
import sys
import tempfile
import time
from collections import Counter, defaultdict
from concurrent.futures import FIRST_COMPLETED, ThreadPoolExecutor, wait
from dataclasses import asdict, dataclass
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any
from zoneinfo import ZoneInfo

try:
    from docx import Document
except Exception:
    Document = None


ROOT = Path(__file__).resolve().parents[1]
PROJECTS_ROOT = ROOT.parent
RESULTS_ROOT = ROOT / "benchmark-results" / "ruflo-repo-optimizer-live"
PROGRESS_PATH = RESULTS_ROOT / "progress.json"
SUMMARY_PATH = RESULTS_ROOT / "summary.json"
RESULTS_PATH = RESULTS_ROOT / "results.jsonl"
RECOMMENDATIONS_PATH = RESULTS_ROOT / "recommended-settings.json"
DOCX_PATH = RESULTS_ROOT / "ruflo-repo-optimizer.docx"
WORKER_STATE_PATH = RESULTS_ROOT / "worker.json"
RUFLO_NORMALIZER = ROOT / "ruflo_config.py"

NPX = shutil.which("npx.cmd") or shutil.which("npx")
NPM = shutil.which("npm.cmd") or shutil.which("npm")
NODE = shutil.which("node.exe") or shutil.which("node")
CLAUDE = shutil.which("claude.cmd") or shutil.which("claude.exe") or shutil.which("claude")
PYTHON = sys.executable

EXCLUDED_REPOS = {
    ".claude-flow",
    "benchmark-results",
    "claude-flow-upstream",
    "ruflo-upstream",
    "swarm-town",
}

PROFILE_JSON_KEYS = {
    "repoName",
    "projectType",
    "primaryLanguage",
    "packageManagers",
    "topLevelDirs",
    "recommendedRunCommand",
}

CONFIG_REPORT_KEYS = {
    "profile",
    "hooksEnabled",
    "statuslineEnabled",
    "autoMemoryEnabled",
    "agentTeamsEnabled",
    "hookKeys",
    "helperFiles",
}

DEFAULT_MAX_PARALLEL = 10
DEFAULT_JOB_TIMEOUT_SECONDS = 300
DEFAULT_WARMUP_CYCLES = 1
DEFAULT_EXPLOIT_VARIANTS = 4
DEFAULT_SCENARIOS_PER_RUN = 2
DEFAULT_COVERAGE_FLOOR = 2
DEFAULT_SCENARIO_FLOOR = 2
DEFAULT_SAMPLE_FLOOR = 4
RUNNER_MODES = ("oneshot", "staged_session")
WORKSPACE_MODES = ("realistic", "compact")
PY_SCENARIO_NAMES = ("py_math_vector", "py_state_machine")
JS_SCENARIO_NAMES = ("js_math_vector", "js_team_allocator")


@dataclass(frozen=True)
class Variant:
    name: str
    profile: str
    hooks: bool = True
    statusline: bool = False
    auto_memory: bool = False
    agent_teams: bool = False


@dataclass(frozen=True)
class RepoSpec:
    name: str
    path: str
    project_type: str
    primary_language: str
    file_count: int


@dataclass(frozen=True)
class RunPlan:
    repo: RepoSpec
    variant: Variant
    runner_mode: str
    workspace_mode: str
    scenarios: tuple[str, ...]
    cycle: int
    job_seed: int


@dataclass(frozen=True)
class ScenarioDef:
    name: str
    language: str
    task_dir: str
    verify_command: tuple[str, ...]
    verify_cwd: str
    objective: str


@dataclass
class CommandResult:
    command: list[str]
    returncode: int
    stdout: str
    stderr: str
    timed_out: bool
    elapsed_ms: int


VARIANTS = [
    Variant("minimal_default", "minimal", hooks=True, statusline=True, auto_memory=False, agent_teams=False),
    Variant("minimal_no_status", "minimal", hooks=True, statusline=False, auto_memory=False, agent_teams=False),
    Variant("minimal_no_hooks", "minimal", hooks=False, statusline=False, auto_memory=False, agent_teams=False),
    Variant("minimal_agent_teams", "minimal", hooks=True, statusline=False, auto_memory=False, agent_teams=True),
    Variant("minimal_status_only", "minimal", hooks=False, statusline=True, auto_memory=False, agent_teams=False),
    Variant("full_default", "full", hooks=True, statusline=True, auto_memory=True, agent_teams=False),
    Variant("full_no_status", "full", hooks=True, statusline=False, auto_memory=True, agent_teams=False),
    Variant("full_no_auto_memory", "full", hooks=True, statusline=False, auto_memory=False, agent_teams=False),
    Variant("full_agent_teams", "full", hooks=True, statusline=False, auto_memory=True, agent_teams=True),
    Variant("full_memory_status", "full", hooks=True, statusline=True, auto_memory=True, agent_teams=True),
]
VARIANT_BY_NAME = {variant.name: variant for variant in VARIANTS}


def utcnow() -> str:
    return datetime.now(timezone.utc).isoformat()


def hidden_process_kwargs(headless: bool = True) -> dict[str, Any]:
    if os.name != "nt" or not headless:
        return {}
    startupinfo = subprocess.STARTUPINFO()
    startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
    startupinfo.wShowWindow = 0
    return {
        "creationflags": subprocess.CREATE_NO_WINDOW,
        "startupinfo": startupinfo,
    }


def run_command(
    command: list[str],
    cwd: Path,
    timeout_seconds: int = 300,
    env: dict[str, str] | None = None,
    *,
    headless: bool = True,
) -> CommandResult:
    started = time.perf_counter()
    proc = subprocess.Popen(
        command,
        cwd=str(cwd),
        env=env if env is not None else os.environ.copy(),
        text=True,
        encoding="utf-8",
        errors="replace",
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        **hidden_process_kwargs(headless=headless),
    )
    try:
        stdout, stderr = proc.communicate(timeout=timeout_seconds)
        return CommandResult(
            command=command,
            returncode=proc.returncode,
            stdout=stdout or "",
            stderr=stderr or "",
            timed_out=False,
            elapsed_ms=int((time.perf_counter() - started) * 1000),
        )
    except subprocess.TimeoutExpired:
        if os.name == "nt":
            subprocess.run(
                ["taskkill", "/PID", str(proc.pid), "/T", "/F"],
                capture_output=True,
                text=True,
                check=False,
                **hidden_process_kwargs(headless=headless),
            )
        else:
            proc.kill()
        stdout, stderr = proc.communicate()
        return CommandResult(
            command=command,
            returncode=-9,
            stdout=stdout or "",
            stderr=(stderr or "").strip() + f"\nProcess timed out after {timeout_seconds} seconds.",
            timed_out=True,
            elapsed_ms=int((time.perf_counter() - started) * 1000),
        )


def run_ruflo(
    args: list[str],
    cwd: Path,
    timeout_seconds: int = 300,
    env: dict[str, str] | None = None,
    *,
    headless: bool = True,
) -> CommandResult:
    if not NPX:
        raise RuntimeError("npx is not available on PATH.")
    return run_command([NPX, "ruflo@latest", *args], cwd=cwd, timeout_seconds=timeout_seconds, env=env, headless=headless)


def normalize_variant(
    project_dir: Path,
    variant: Variant,
    env: dict[str, str] | None = None,
    *,
    headless: bool = True,
) -> CommandResult:
    command = [
        PYTHON,
        str(RUFLO_NORMALIZER),
        "normalize",
        "--project",
        str(project_dir),
        "--profile",
        variant.profile,
    ]
    if not variant.hooks:
        command.append("--disable-hooks")
    if not variant.statusline:
        command.append("--disable-statusline")
    if not variant.auto_memory:
        command.append("--disable-auto-memory")
    if variant.agent_teams:
        command.append("--agent-teams")
    return run_command(command, cwd=ROOT, timeout_seconds=180, env=env, headless=headless)


def ignore_filter(_: str, names: list[str]) -> set[str]:
    ignored = set()
    for name in names:
        if name in {".git", ".hg", ".svn", "node_modules", "dist", "build", "coverage", ".next", ".nuxt", ".cache", ".turbo", ".venv", "venv", "__pycache__", ".claude", ".claude-flow", ".swarm"}:
            ignored.add(name)
        elif name.upper() == "NUL":
            ignored.add(name)
        elif name.endswith((".log", ".sqlite", ".db-wal", ".db-shm")):
            ignored.add(name)
    return ignored


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8", newline="\n")


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def read_json(path: Path, default: Any) -> Any:
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return default


def load_json(path: Path, default: Any) -> Any:
    return read_json(path, default)


def detect_repo(path: Path) -> RepoSpec | None:
    files = {entry.name.lower() for entry in path.iterdir()} if path.is_dir() else set()
    if not files:
        return None

    js_files = 0
    py_files = 0
    for root, dirs, fnames in os.walk(path):
        dirs[:] = [d for d in dirs if d not in {".git", "node_modules", "__pycache__", ".next", "venv", "dist", "build", ".claude", ".claude-flow"}]
        for fname in fnames:
            if fname.endswith((".js", ".jsx", ".ts", ".tsx")):
                js_files += 1
            elif fname.endswith(".py"):
                py_files += 1

    project_type = "unknown"
    primary_language = "unknown"

    if "package.json" in files:
        project_type = "node"
        primary_language = "TypeScript" if any(name.endswith((".ts", ".tsx")) for name in files) or js_files else "JavaScript"
        try:
            pkg = json.loads((path / "package.json").read_text(encoding="utf-8"))
            deps = {**pkg.get("dependencies", {}), **pkg.get("devDependencies", {})}
            if "react" in deps or "next" in deps:
                project_type = "node-react"
            if "typescript" in deps:
                primary_language = "TypeScript"
        except Exception:
            pass
    if "pyproject.toml" in files or "requirements.txt" in files or py_files:
        if project_type != "unknown":
            project_type = "fullstack"
        else:
            project_type = "python"
        if primary_language == "unknown":
            primary_language = "Python"
    if "cargo.toml" in files:
        project_type = "rust"
        primary_language = "Rust"
    if project_type == "unknown" and any(name.endswith(".html") for name in files):
        project_type = "static"
        primary_language = "HTML"

    file_count = js_files + py_files
    if project_type == "unknown" and file_count == 0:
        return None

    return RepoSpec(
        name=path.name,
        path=str(path),
        project_type=project_type,
        primary_language=primary_language,
        file_count=file_count,
    )


def discover_repos(projects_root: Path, repo_names: list[str] | None, max_repos: int) -> list[RepoSpec]:
    repos: list[RepoSpec] = []
    candidates = [projects_root / name for name in repo_names] if repo_names else sorted(projects_root.iterdir())
    for candidate in candidates:
        if not candidate.is_dir():
            continue
        if candidate.name in EXCLUDED_REPOS or candidate.name.startswith("."):
            continue
        spec = detect_repo(candidate)
        if spec:
            repos.append(spec)
        if len(repos) >= max_repos:
            break
    return repos


def compact_copy_repo(repo: RepoSpec, destination: Path) -> None:
    source = Path(repo.path)
    destination.mkdir(parents=True, exist_ok=True)
    copied = 0
    max_files = 450
    max_bytes = 1_500_000
    for root, dirs, files in os.walk(source):
        rel_root = Path(root).relative_to(source)
        if len(rel_root.parts) > 5:
            dirs[:] = []
            continue
        dirs[:] = [
            item
            for item in dirs
            if item not in {".git", "node_modules", "__pycache__", ".next", "dist", "build", ".cache", ".venv", "venv"}
        ]
        if rel_root.parts and rel_root.parts[0].startswith(".") and rel_root.parts[0] not in {".claude", ".claude-flow"}:
            dirs[:] = []
            continue
        for fname in files:
            if copied >= max_files:
                return
            if fname.endswith((".pyc", ".log", ".db", ".sqlite", ".tmp")):
                continue
            src = Path(root) / fname
            try:
                size = src.stat().st_size
            except OSError:
                continue
            if size > max_bytes:
                continue
            rel = src.relative_to(source)
            if rel.parts and rel.parts[0].startswith(".") and rel.parts[0] not in {".claude", ".claude-flow"}:
                continue
            dst = destination / rel
            dst.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(src, dst)
            copied += 1
    for root_name in ("package.json", "pyproject.toml", "requirements.txt", "README.md"):
        src = source / root_name
        if src.exists():
            dst = destination / root_name
            if not dst.exists():
                shutil.copy2(src, dst)


def copy_repo(repo: RepoSpec, destination: Path, workspace_mode: str = "realistic") -> None:
    if workspace_mode == "compact":
        compact_copy_repo(repo, destination)
        return
    shutil.copytree(Path(repo.path), destination, ignore=ignore_filter, dirs_exist_ok=True)


def build_isolated_env(workspace: Path, run_key: str) -> dict[str, str]:
    env = os.environ.copy()
    cache_root = workspace / ".ruflo-benchmark" / ".cache" / run_key
    xdg = cache_root / "xdg"
    npm_cache = cache_root / "npm"
    pip_cache = cache_root / "pip"
    pycache = cache_root / "pycache"
    for path in (xdg, npm_cache, pip_cache, pycache):
        path.mkdir(parents=True, exist_ok=True)
    env["XDG_CACHE_HOME"] = str(xdg)
    env["NPM_CONFIG_CACHE"] = str(npm_cache)
    env["PIP_CACHE_DIR"] = str(pip_cache)
    env["PYTHONPYCACHEPREFIX"] = str(pycache)
    env["PYTHONDONTWRITEBYTECODE"] = "1"
    env["CLAUDE_PROJECT_DIR"] = str(workspace)
    env["NO_COLOR"] = "1"
    return env


def package_managers_for_repo(root: Path) -> list[str]:
    managers: list[str] = []
    nested_package = next((path for path in root.glob("*/package.json")), None)
    nested_pyproject = next((path for path in root.glob("*/pyproject.toml")), None)
    nested_requirements = next((path for path in root.glob("*/requirements.txt")), None)
    if (root / "package.json").exists() or nested_package:
        managers.append("npm")
    if (root / "pnpm-lock.yaml").exists():
        managers.append("pnpm")
    if (root / "yarn.lock").exists():
        managers.append("yarn")
    if (root / "requirements.txt").exists() or (root / "pyproject.toml").exists() or nested_pyproject or nested_requirements:
        managers.append("pip")
    if (root / "Cargo.toml").exists():
        managers.append("cargo")
    return managers or ["unknown"]


def guess_run_command(root: Path, repo: RepoSpec) -> str:
    package_json = root / "package.json"
    nested_package = next((path for path in root.glob("*/package.json")), None)
    if package_json.exists():
        try:
            pkg = json.loads(package_json.read_text(encoding="utf-8"))
            scripts = pkg.get("scripts") or {}
            for key in ("dev", "start", "test"):
                if key in scripts:
                    return f"npm run {key}"
        except Exception:
            pass
        return "npm run dev"
    if nested_package:
        try:
            pkg = json.loads(nested_package.read_text(encoding="utf-8"))
            scripts = pkg.get("scripts") or {}
            for key in ("dev", "start", "test"):
                if key in scripts:
                    return f"cd {nested_package.parent.name} && npm run {key}"
        except Exception:
            pass
        return f"cd {nested_package.parent.name} && npm run dev"
    if (root / "manage.py").exists():
        return "python manage.py runserver"
    if (root / "app.py").exists():
        return "python app.py"
    if (root / "main.py").exists():
        return "python main.py"
    if repo.primary_language == "Python":
        return "python -m unittest discover"
    return "npm test"


def write_repo_context(root: Path, repo: RepoSpec) -> dict[str, Any]:
    top_level = []
    for entry in sorted(root.iterdir(), key=lambda item: item.name.lower()):
        if entry.name.startswith(".") and entry.name not in {".claude", ".claude-flow", ".ruflo-benchmark"}:
            continue
        if entry.name in {"node_modules", "__pycache__", "dist", "build"}:
            continue
        top_level.append(entry.name)
        if len(top_level) >= 12:
            break
    context = {
        "repoName": repo.name,
        "projectType": repo.project_type,
        "primaryLanguage": repo.primary_language,
        "packageManagers": package_managers_for_repo(root),
        "topLevelDirs": top_level,
        "recommendedRunCommand": guess_run_command(root, repo),
    }
    return context


def config_expectations(root: Path) -> dict[str, Any]:
    settings_path = root / ".claude" / "settings.json"
    runtime_path = root / ".claude-flow" / "config.json"
    helpers_path = root / ".claude" / "helpers"
    settings = json.loads(settings_path.read_text(encoding="utf-8")) if settings_path.exists() else {}
    runtime = json.loads(runtime_path.read_text(encoding="utf-8")) if runtime_path.exists() else {}
    hooks = settings.get("hooks") or {}
    features = runtime.get("features") or {}
    helper_files = sorted(
        path.name
        for path in helpers_path.iterdir()
        if path.is_file()
    ) if helpers_path.exists() else []
    return {
        "profile": runtime.get("profile"),
        "hooksEnabled": bool(features.get("hooks")) or bool(hooks),
        "statuslineEnabled": bool(features.get("statusLine")) or bool(settings.get("statusLine")),
        "autoMemoryEnabled": bool(features.get("autoMemory")),
        "agentTeamsEnabled": bool(features.get("agentTeams")),
        "hookKeys": sorted(hooks.keys()) if isinstance(hooks, dict) else [],
        "helperFiles": helper_files,
    }




@dataclass(frozen=True)
class ModeSpec:
    name: str
    staged: bool


@dataclass(frozen=True)
class ScenarioDefinition:
    name: str
    family: str
    language: str
    description: str


@dataclass(frozen=True)
class ScenarioPlan:
    definition: ScenarioDefinition
    focus_paths: tuple[str, ...]
    verify_commands: tuple[tuple[str, ...], ...]
    created_files: tuple[str, ...]
    prompt_fragment: str


RUNNER_MODES = {
    'oneshot': ModeSpec(name='oneshot', staged=False),
    'staged_session': ModeSpec(name='staged_session', staged=True),
}
DEFAULT_RUNNER_MODES = tuple(RUNNER_MODES.keys())
DEFAULT_WORKSPACE_MODE = 'realistic'
GLOBAL_MIN_SAMPLES = 24
PROJECT_TYPE_MIN_SAMPLES = 12
REPO_MIN_SAMPLES = 6
EXPLORATION_FLOOR = 2
PROMPT_STYLE_VARIANTS = ('direct', 'structured', 'constraint_first')
PROMPT_STYLE_TEXT = {
    'direct': 'Work directly and keep edits minimal.',
    'structured': 'Reason from tests to code paths, then make the smallest coherent fix set.',
    'constraint_first': 'Honor the benchmark constraints first, then fix correctness with explicit verification.',
}
SCENARIOS = {
    'js_multifile_sum': ScenarioDefinition(
        name='js_multifile_sum',
        family='js_multifile',
        language='javascript',
        description='Fix a multi-file arithmetic flow in common Node-style source and test locations.',
    ),
    'js_manifest_renderer': ScenarioDefinition(
        name='js_manifest_renderer',
        family='js_rendering',
        language='javascript',
        description='Fix ordering and formatting across coupled manifest/render modules.',
    ),
    'py_multifile_math': ScenarioDefinition(
        name='py_multifile_math',
        family='py_multifile',
        language='python',
        description='Fix a Python helper plus consumer module pair with a failing test.',
    ),
    'py_report_builder': ScenarioDefinition(
        name='py_report_builder',
        family='py_rendering',
        language='python',
        description='Fix parse/render behavior spread across multiple Python modules.',
    ),
}


def deterministic_int(*parts: Any) -> int:
    payload = '||'.join(str(part) for part in parts)
    digest = hashlib.sha256(payload.encode('utf-8')).hexdigest()
    return int(digest[:16], 16)


def relative_string(root: Path, path: Path) -> str:
    return path.relative_to(root).as_posix()


def choose_source_dir(root: Path) -> Path:
    preferred = [
        root / 'src',
        root / 'lib',
        root / 'app',
        root / 'server',
        root / 'client' / 'src',
        root / 'packages' / 'app' / 'src',
    ]
    for candidate in preferred:
        if candidate.exists() and candidate.is_dir():
            candidate.mkdir(parents=True, exist_ok=True)
            return candidate
    package_dirs = sorted(
        path.parent
        for path in root.rglob('package.json')
        if '.claude' not in path.parts and '.claude-flow' not in path.parts and 'node_modules' not in path.parts
    )
    for pkg_dir in package_dirs:
        for child_name in ('src', 'lib', 'app'):
            candidate = pkg_dir / child_name
            if candidate.exists() and candidate.is_dir():
                candidate.mkdir(parents=True, exist_ok=True)
                return candidate
        return pkg_dir
    python_files = sorted(
        path.parent
        for path in root.rglob('*.py')
        if '.claude' not in path.parts and '.claude-flow' not in path.parts and '__pycache__' not in path.parts
    )
    if python_files:
        return python_files[0]
    fallback = root / 'src'
    fallback.mkdir(parents=True, exist_ok=True)
    return fallback


def choose_test_dir(root: Path) -> Path:
    preferred = [
        root / 'tests',
        root / 'test',
        root / '__tests__',
        root / 'spec',
        root / 'src' / '__tests__',
    ]
    for candidate in preferred:
        if candidate.exists() and candidate.is_dir():
            candidate.mkdir(parents=True, exist_ok=True)
            return candidate
    package_dirs = sorted(
        path.parent
        for path in root.rglob('package.json')
        if '.claude' not in path.parts and '.claude-flow' not in path.parts and 'node_modules' not in path.parts
    )
    for pkg_dir in package_dirs:
        for child_name in ('tests', 'test', '__tests__', 'spec'):
            candidate = pkg_dir / child_name
            if candidate.exists() and candidate.is_dir():
                candidate.mkdir(parents=True, exist_ok=True)
                return candidate
    fallback = root / 'tests'
    fallback.mkdir(parents=True, exist_ok=True)
    return fallback


def summarize_workspace(root: Path, workspace_mode: str) -> dict[str, Any]:
    total_files = 0
    benchmark_files = 0
    top_level: list[str] = []
    for entry in sorted(root.iterdir(), key=lambda item: item.name.lower()):
        if entry.name.startswith('.') and entry.name not in {'.claude', '.claude-flow', '.ruflo-benchmark'}:
            continue
        top_level.append(entry.name)
        if len(top_level) >= 12:
            break
    for path in root.rglob('*'):
        if path.is_file():
            total_files += 1
            if '.ruflo-benchmark' in path.parts:
                benchmark_files += 1
    return {
        'mode': workspace_mode,
        'totalFiles': total_files,
        'benchmarkFiles': benchmark_files,
        'topLevelEntries': top_level,
    }


def compact_copy_repo(repo: RepoSpec, destination: Path) -> dict[str, Any]:
    source = Path(repo.path)
    destination.mkdir(parents=True, exist_ok=True)
    copied = 0
    max_files = 650
    max_bytes = 2_500_000
    for root_dir, dirs, files in os.walk(source):
        rel_root = Path(root_dir).relative_to(source)
        if len(rel_root.parts) > 6:
            dirs[:] = []
            continue
        dirs[:] = [
            item
            for item in dirs
            if item not in {'.git', 'node_modules', '__pycache__', '.next', 'dist', 'build', '.cache', '.venv', 'venv'}
        ]
        if rel_root.parts and rel_root.parts[0].startswith('.') and rel_root.parts[0] not in {'.claude', '.claude-flow'}:
            dirs[:] = []
            continue
        for fname in files:
            if copied >= max_files:
                return summarize_workspace(destination, 'compact')
            if fname.endswith(('.pyc', '.log', '.db', '.sqlite', '.tmp')):
                continue
            src = Path(root_dir) / fname
            try:
                size = src.stat().st_size
            except OSError:
                continue
            if size > max_bytes:
                continue
            rel = src.relative_to(source)
            if rel.parts and rel.parts[0].startswith('.') and rel.parts[0] not in {'.claude', '.claude-flow'}:
                continue
            dst = destination / rel
            dst.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(src, dst)
            copied += 1
    for root_name in ('package.json', 'pyproject.toml', 'requirements.txt', 'README.md'):
        src = source / root_name
        if src.exists():
            dst = destination / root_name
            if not dst.exists():
                shutil.copy2(src, dst)
    return summarize_workspace(destination, 'compact')


def copy_repo(repo: RepoSpec, destination: Path, workspace_mode: str = DEFAULT_WORKSPACE_MODE) -> dict[str, Any]:
    if workspace_mode == 'compact':
        return compact_copy_repo(repo, destination)
    shutil.copytree(Path(repo.path), destination, ignore=ignore_filter, dirs_exist_ok=True)
    return summarize_workspace(destination, 'realistic')


def write_repo_context(root: Path, repo: RepoSpec, workspace_profile: dict[str, Any] | None = None) -> dict[str, Any]:
    top_level = []
    for entry in sorted(root.iterdir(), key=lambda item: item.name.lower()):
        if entry.name.startswith('.') and entry.name not in {'.claude', '.claude-flow', '.ruflo-benchmark'}:
            continue
        if entry.name in {'node_modules', '__pycache__', 'dist', 'build'}:
            continue
        top_level.append(entry.name)
        if len(top_level) >= 12:
            break
    context = {
        'repoName': repo.name,
        'projectType': repo.project_type,
        'primaryLanguage': repo.primary_language,
        'packageManagers': package_managers_for_repo(root),
        'topLevelDirs': top_level,
        'recommendedRunCommand': guess_run_command(root, repo),
    }
    if workspace_profile is not None:
        context['workspaceProfile'] = workspace_profile
    write_json(root / '.ruflo-benchmark' / 'gold' / 'repo-context.json', context)
    return context


def run_shell_command(command: str, cwd: Path, timeout_seconds: int = 30, input_text: str | None = None, *, headless: bool = True) -> CommandResult:
    started = time.perf_counter()
    proc = subprocess.Popen(
        command,
        cwd=str(cwd),
        env=os.environ.copy(),
        text=True,
        encoding='utf-8',
        errors='replace',
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        shell=True,
        **hidden_process_kwargs(headless=headless),
    )
    try:
        stdout, stderr = proc.communicate(input=input_text, timeout=timeout_seconds)
        return CommandResult(
            command=[command],
            returncode=proc.returncode,
            stdout=stdout or '',
            stderr=stderr or '',
            timed_out=False,
            elapsed_ms=int((time.perf_counter() - started) * 1000),
        )
    except subprocess.TimeoutExpired:
        if os.name == 'nt':
            subprocess.run(['taskkill', '/PID', str(proc.pid), '/T', '/F'], capture_output=True, text=True, check=False, **hidden_process_kwargs(headless=headless))
        else:
            proc.kill()
        stdout, stderr = proc.communicate()
        return CommandResult(
            command=[command],
            returncode=-9,
            stdout=stdout or '',
            stderr=(stderr or '').strip() + f"\\nProcess timed out after {timeout_seconds} seconds.",
            timed_out=True,
            elapsed_ms=int((time.perf_counter() - started) * 1000),
        )
def eligible_scenarios(repo: RepoSpec) -> list[ScenarioDefinition]:
    if repo.primary_language == "Python" or repo.project_type == "python":
        return [SCENARIOS["py_multifile_math"], SCENARIOS["py_report_builder"]]
    return [SCENARIOS["js_multifile_sum"], SCENARIOS["js_manifest_renderer"]]


def build_js_multifile_sum(root: Path) -> ScenarioPlan:
    source_dir = choose_source_dir(root)
    test_dir = choose_test_dir(root)
    task_dir = source_dir / "ruflo_bench_sum"
    test_path = test_dir / "ruflo_bench_sum.test.mjs"
    math_path = task_dir / "math.mjs"
    summary_path = task_dir / "summary.mjs"
    write_text(math_path, "export function addNumbers(values) {\n  return values.reduce((total, value) => total - value, 0);\n}\n")
    write_text(
        summary_path,
        "import { addNumbers } from './math.mjs';\n"
        "export function summarizeValues(values) {\n"
        "  return {\n"
        "    count: values.length,\n"
        "    total: addNumbers(values),\n"
        "    average: values.length ? addNumbers(values) / values.length : 0,\n"
        "  };\n"
        "}\n",
    )
    import_summary = os.path.relpath(summary_path, test_dir).replace("\\", "/")
    write_text(
        test_path,
        "import test from 'node:test';\n"
        "import assert from 'node:assert/strict';\n"
        f"import {{ summarizeValues }} from './{import_summary}';\n\n"
        "test('summarizeValues returns correct totals', () => {\n"
        "  assert.deepEqual(summarizeValues([2, 3, 4]), { count: 3, total: 9, average: 3 });\n"
        "});\n",
    )
    return ScenarioPlan(
        definition=SCENARIOS["js_multifile_sum"],
        focus_paths=(relative_string(root, math_path), relative_string(root, summary_path), relative_string(root, test_path)),
        verify_commands=((NODE or "node", "--test", relative_string(root, test_path)),),
        created_files=(relative_string(root, math_path), relative_string(root, summary_path), relative_string(root, test_path)),
        prompt_fragment=f"Fix the multi-file benchmark under {relative_string(root, task_dir)} so {relative_string(root, test_path)} passes. The defect may span more than one file.",
    )


def build_js_manifest_renderer(root: Path) -> ScenarioPlan:
    source_dir = choose_source_dir(root)
    test_dir = choose_test_dir(root)
    task_dir = source_dir / "ruflo_bench_manifest"
    manifest_path = task_dir / "manifest.mjs"
    render_path = task_dir / "render.mjs"
    test_path = test_dir / "ruflo_bench_manifest.test.mjs"
    write_text(
        manifest_path,
        "export function buildManifest(entries) {\n"
        "  return entries.map((entry) => ({\n"
        "    slug: entry.slug,\n"
        "    title: entry.title.toLowerCase(),\n"
        "    words: entry.words,\n"
        "  }));\n"
        "}\n",
    )
    write_text(
        render_path,
        "import { buildManifest } from './manifest.mjs';\n"
        "export function renderManifest(entries) {\n"
        "  return buildManifest(entries)\n"
        "    .sort((left, right) => right.slug.localeCompare(left.slug))\n"
        "    .map((entry) => `${entry.slug}:${entry.title}:${entry.words}`)\n"
        "    .join('\\n');\n"
        "}\n",
    )
    import_render = os.path.relpath(render_path, test_dir).replace("\\", "/")
    write_text(
        test_path,
        "import test from 'node:test';\n"
        "import assert from 'node:assert/strict';\n"
        f"import {{ renderManifest }} from './{import_render}';\n\n"
        "test('renderManifest preserves title casing and ascending slug order', () => {\n"
        "  const result = renderManifest([\n"
        "    { slug: 'beta', title: 'Beta Story', words: 120 },\n"
        "    { slug: 'alpha', title: 'Alpha Story', words: 95 },\n"
        "  ]);\n"
        "  assert.equal(result, 'alpha:Alpha Story:95\\nbeta:Beta Story:120');\n"
        "});\n",
    )
    return ScenarioPlan(
        definition=SCENARIOS["js_manifest_renderer"],
        focus_paths=(relative_string(root, manifest_path), relative_string(root, render_path), relative_string(root, test_path)),
        verify_commands=((NODE or "node", "--test", relative_string(root, test_path)),),
        created_files=(relative_string(root, manifest_path), relative_string(root, render_path), relative_string(root, test_path)),
        prompt_fragment=f"Fix the manifest/render benchmark in {relative_string(root, task_dir)} so {relative_string(root, test_path)} passes. Preserve the intended output format.",
    )


def build_py_multifile_math(root: Path) -> ScenarioPlan:
    source_dir = choose_source_dir(root)
    test_dir = choose_test_dir(root)
    task_dir = source_dir / "ruflo_bench_math"
    helpers_path = task_dir / "helpers.py"
    calc_path = task_dir / "calc.py"
    test_path = test_dir / "test_ruflo_bench_math.py"
    write_text(helpers_path, "def sum_values(values):\n    total = 0\n    for value in values:\n        total -= value\n    return total\n")
    write_text(
        calc_path,
        "from helpers import sum_values\n\n"
        "def summarize(values):\n"
        "    total = sum_values(values)\n"
        "    return {\n"
        "        'count': len(values),\n"
        "        'total': total,\n"
        "        'average': total / len(values) if values else 0,\n"
        "    }\n",
    )
    calc_import_dir = calc_path.parent
    write_text(
        test_path,
        "import sys\nimport unittest\n\n"
        f"sys.path.insert(0, r'{calc_import_dir}')\n"
        "from calc import summarize\n\n"
        "class BenchMathTest(unittest.TestCase):\n"
        "    def test_summarize(self):\n"
        "        self.assertEqual(summarize([2, 3, 4]), {'count': 3, 'total': 9, 'average': 3})\n\n"
        "if __name__ == '__main__':\n"
        "    unittest.main()\n",
    )
    return ScenarioPlan(
        definition=SCENARIOS["py_multifile_math"],
        focus_paths=(relative_string(root, helpers_path), relative_string(root, calc_path), relative_string(root, test_path)),
        verify_commands=((PYTHON, relative_string(root, test_path)),),
        created_files=(relative_string(root, helpers_path), relative_string(root, calc_path), relative_string(root, test_path)),
        prompt_fragment=f"Fix the Python benchmark under {relative_string(root, task_dir)} so {relative_string(root, test_path)} passes. The bug may be split across helper and consumer modules.",
    )


def build_py_report_builder(root: Path) -> ScenarioPlan:
    source_dir = choose_source_dir(root)
    test_dir = choose_test_dir(root)
    task_dir = source_dir / "ruflo_bench_report"
    parse_path = task_dir / "parse_report.py"
    render_path = task_dir / "render_report.py"
    test_path = test_dir / "test_ruflo_bench_report.py"
    write_text(parse_path, "def parse_rows(rows):\n    parsed = []\n    for slug, title, words in rows:\n        parsed.append({'slug': slug, 'title': title.lower(), 'words': words})\n    return parsed\n")
    write_text(
        render_path,
        "from parse_report import parse_rows\n\n"
        "def render_rows(rows):\n"
        "    parsed = parse_rows(rows)\n"
        "    parsed.sort(key=lambda item: item['slug'], reverse=True)\n"
        "    return '\\n'.join(f\"{item['slug']}:{item['title']}:{item['words']}\" for item in parsed)\n",
    )
    render_import_dir = render_path.parent
    write_text(
        test_path,
        "import sys\nimport unittest\n\n"
        f"sys.path.insert(0, r'{render_import_dir}')\n"
        "from render_report import render_rows\n\n"
        "class BenchReportTest(unittest.TestCase):\n"
        "    def test_render_rows(self):\n"
        "        rows = [('beta', 'Beta Story', 120), ('alpha', 'Alpha Story', 95)]\n"
        "        self.assertEqual(render_rows(rows), 'alpha:Alpha Story:95\\nbeta:Beta Story:120')\n\n"
        "if __name__ == '__main__':\n"
        "    unittest.main()\n",
    )
    return ScenarioPlan(
        definition=SCENARIOS["py_report_builder"],
        focus_paths=(relative_string(root, parse_path), relative_string(root, render_path), relative_string(root, test_path)),
        verify_commands=((PYTHON, relative_string(root, test_path)),),
        created_files=(relative_string(root, parse_path), relative_string(root, render_path), relative_string(root, test_path)),
        prompt_fragment=f"Fix the Python report benchmark under {relative_string(root, task_dir)} so {relative_string(root, test_path)} passes. Preserve title casing and ascending slug order.",
    )


SCENARIO_BUILDERS = {
    "js_multifile_sum": build_js_multifile_sum,
    "js_manifest_renderer": build_js_manifest_renderer,
    "py_multifile_math": build_py_multifile_math,
    "py_report_builder": build_py_report_builder,
}


def scenario_history_counts(results: list[dict[str, Any]], repo: RepoSpec, runner_mode: str) -> Counter[str]:
    return Counter(row["scenario"] for row in results if row.get("repo") == repo.name and row.get("runnerMode") == runner_mode and row.get("scenario"))


def choose_runner_mode(repo: RepoSpec, variant: Variant, results: list[dict[str, Any]], runner_modes: list[str], cycle: int) -> str:
    variant_counts = Counter(
        row["runnerMode"]
        for row in results
        if row.get("repo") == repo.name and row.get("variant") == variant.name and row.get("runnerMode") in runner_modes
    )
    repo_counts = Counter(
        row["runnerMode"]
        for row in results
        if row.get("repo") == repo.name and row.get("runnerMode") in runner_modes
    )
    ordered = sorted(
        runner_modes,
        key=lambda name: (
            variant_counts.get(name, 0),
            repo_counts.get(name, 0),
            deterministic_int(repo.name, variant.name, cycle, name),
        ),
    )
    return ordered[0]


def choose_prompt_style(repo: RepoSpec, scenario_name: str, runner_mode: str, results: list[dict[str, Any]], cycle: int) -> str:
    counts = Counter(row["promptStyle"] for row in results if row.get("repo") == repo.name and row.get("scenario") == scenario_name and row.get("runnerMode") == runner_mode)
    ordered = sorted(PROMPT_STYLE_VARIANTS, key=lambda name: (counts.get(name, 0), deterministic_int(repo.name, scenario_name, runner_mode, name, cycle)))
    return ordered[0]


def choose_scenario(repo: RepoSpec, runner_mode: str, results: list[dict[str, Any]], cycle: int) -> ScenarioDefinition:
    eligible = eligible_scenarios(repo)
    counts = scenario_history_counts(results, repo, runner_mode)
    ordered = sorted(eligible, key=lambda item: (counts.get(item.name, 0), deterministic_int(repo.name, runner_mode, item.name, cycle)))
    return ordered[0]


def build_scenario_plan(root: Path, repo: RepoSpec, scenario_name: str) -> ScenarioPlan:
    return SCENARIO_BUILDERS[scenario_name](root)


def remaining_timeout(started_at: float, total_seconds: int, stage_cap: int, reserve_seconds: int = 0) -> int:
    elapsed = time.perf_counter() - started_at
    remaining = int(total_seconds - elapsed - reserve_seconds)
    return max(1, min(stage_cap, remaining))


def claude_command(prompt: str, effort: str = "low", session_persistence: bool = False) -> list[str]:
    if not CLAUDE:
        raise RuntimeError("Claude Code CLI not found on PATH.")
    command = [
        CLAUDE,
        "-p",
        prompt,
        "--model",
        "sonnet",
        "--effort",
        effort,
        "--output-format",
        "json",
        "--permission-mode",
        "bypassPermissions",
        "--dangerously-skip-permissions",
        "--allowedTools",
        "Bash",
        "Read",
        "Write",
        "Edit",
        "Glob",
        "Grep",
        "LS",
    ]
    if not session_persistence:
        command.append("--no-session-persistence")
    return command


def parse_claude(stdout: str) -> dict[str, Any]:
    try:
        return json.loads(stdout.strip()) if stdout.strip() else {}
    except json.JSONDecodeError as exc:
        return {"parse_error": str(exc), "raw": stdout}


def aggregate_claude_runs(results: list[tuple[CommandResult, dict[str, Any], str]]) -> tuple[CommandResult, dict[str, Any]]:
    combined_stdout = "\n".join(result.stdout for result, _, _ in results if result.stdout)
    combined_stderr = "\n".join(result.stderr for result, _, _ in results if result.stderr)
    total_duration_ms = sum(int((parsed.get("duration_ms") or 0)) for _, parsed, _ in results)
    total_cost = sum(float(parsed.get("total_cost_usd") or 0) for _, parsed, _ in results)
    total_input = sum(int(((parsed.get("usage") or {}).get("input_tokens") or 0)) for _, parsed, _ in results)
    total_output = sum(int(((parsed.get("usage") or {}).get("output_tokens") or 0)) for _, parsed, _ in results)
    combined = {
        "duration_ms": total_duration_ms,
        "total_cost_usd": round(total_cost, 6),
        "usage": {"input_tokens": total_input, "output_tokens": total_output},
        "steps": [
            {
                "name": step_name,
                "duration_ms": parsed.get("duration_ms"),
                "total_cost_usd": parsed.get("total_cost_usd"),
                "result": parsed.get("result"),
                "session_id": parsed.get("session_id"),
            }
            for _, parsed, step_name in results
        ],
        "result": results[-1][1].get("result") if results else None,
    }
    command_repr = " && ".join(" ".join(result.command) if isinstance(result.command, list) else str(result.command) for result, _, _ in results)
    aggregate_result = CommandResult(command=command_repr, returncode=results[-1][0].returncode if results else 0, stdout=combined_stdout, stderr=combined_stderr, timed_out=any(result.timed_out for result, _, _ in results), elapsed_ms=sum(result.elapsed_ms for result, _, _ in results))
    return aggregate_result, combined


def next_anchorage_reset(now: datetime | None = None) -> datetime:
    anchorage = ZoneInfo("America/Anchorage")
    local_now = (now or datetime.now(timezone.utc)).astimezone(anchorage)
    next_midnight = (local_now + timedelta(days=1)).replace(hour=0, minute=0, second=0, microsecond=0)
    return next_midnight.astimezone(timezone.utc)


def is_limit_exceeded(parsed: dict[str, Any], stdout: str, stderr: str) -> bool:
    haystacks = [str(parsed.get("result") or ""), str(parsed.get("raw") or ""), stdout or "", stderr or ""]
    return any("hit your limit" in text.lower() for text in haystacks)


def build_blind_profile_prompt() -> str:
    return (
        "Without using any ground-truth metadata file, inspect the repository and write "
        ".ruflo-benchmark/repo-profile-blind.json with keys repoName, projectType, primaryLanguage, "
        "packageManagers, topLevelDirs, recommendedRunCommand. Ignore benchmark-created files and generated "
        "Claude/Ruflo config when inferring project identity. Prefer the root folder name for repoName."
    )


def build_grounded_profile_prompt() -> str:
    return "After the blind profile is written, open .ruflo-benchmark/gold/repo-context.json and write .ruflo-benchmark/repo-profile-grounded.json with the exact canonical values from that file."


def build_config_report_prompt() -> str:
    return "Inspect .claude/settings.json, .claude-flow/config.json, and .claude/helpers. Write .ruflo-benchmark/config-report.json with keys profile, hooksEnabled, statuslineEnabled, autoMemoryEnabled, agentTeamsEnabled, hookKeys, helperFiles. Use helperFiles as a sorted list of filenames only."


def build_task_report_prompt() -> str:
    return "Write .ruflo-benchmark/task-report.json with keys scenarioName, filesEdited, promptStyle, runnerMode, notes. Keep notes short."


def build_main_prompt(repo: RepoSpec, scenario: ScenarioPlan, runner_mode: str, prompt_style: str, run_command_hint: str) -> str:
    focus = ", ".join(scenario.focus_paths)
    created = ", ".join(scenario.created_files)
    return " ".join([
        PROMPT_STYLE_TEXT[prompt_style],
        "You may edit the repo workspace and the benchmark-created files, but do not modify .claude/settings.json, .claude-flow/config.json, or .ruflo-benchmark/gold/repo-context.json.",
        build_blind_profile_prompt(),
        build_grounded_profile_prompt(),
        build_config_report_prompt(),
        scenario.prompt_fragment,
        f"Focus first on these files: {focus}.",
        f"Benchmark-created files for this scenario: {created}.",
        f"If you need a run command hint, start from: {run_command_hint}.",
        build_task_report_prompt(),
        f"Record runnerMode={runner_mode} and promptStyle={prompt_style} in task-report.json.",
        "Reply only with TASK_DONE.",
    ])


def build_warmup_prompt(repo: RepoSpec, prompt_style: str) -> str:
    return " ".join([
        PROMPT_STYLE_TEXT[prompt_style],
        "Inspect only the top-level repo shape, package metadata, and generated Ruflo/Claude config. Do not edit the benchmark task files yet.",
        "Write .ruflo-benchmark/session-notes.json with keys repoSummary, configSummary, likelyFixPlan, likelyRunCommand.",
        f"Use repoName={repo.name} only as an external label; derive all other details from the workspace.",
        "Reply only with SESSION_READY.",
    ])


def run_claude_flow(workspace: Path, repo: RepoSpec, scenario: ScenarioPlan, runner_mode: str, prompt_style: str, run_command_hint: str, timeout_seconds: int) -> tuple[CommandResult, dict[str, Any], bool]:
    mode = RUNNER_MODES[runner_mode]
    runs: list[tuple[CommandResult, dict[str, Any], str]] = []
    limit_exceeded = False

    if mode.staged:
        warmup_timeout = max(12, min(25, timeout_seconds // 3))
        warmup = run_command(
            claude_command(build_warmup_prompt(repo, prompt_style), effort="low", session_persistence=False),
            cwd=workspace,
            timeout_seconds=warmup_timeout,
        )
        warmup_parsed = parse_claude(warmup.stdout)
        limit_exceeded = is_limit_exceeded(warmup_parsed, warmup.stdout, warmup.stderr)
        if limit_exceeded:
            runs.append((warmup, warmup_parsed, "warmup"))
            aggregate_result, aggregate_parsed = aggregate_claude_runs(runs)
            return aggregate_result, aggregate_parsed, limit_exceeded
        if warmup.returncode == 0 and not warmup.timed_out:
            runs.append((warmup, warmup_parsed, "warmup"))

    main_result = run_command(claude_command(build_main_prompt(repo, scenario, runner_mode, prompt_style, run_command_hint), effort="medium" if mode.staged else "low", session_persistence=False), cwd=workspace, timeout_seconds=max(20, timeout_seconds))
    main_parsed = parse_claude(main_result.stdout)
    runs.append((main_result, main_parsed, "task"))
    limit_exceeded = limit_exceeded or is_limit_exceeded(main_parsed, main_result.stdout, main_result.stderr)
    aggregate_result, aggregate_parsed = aggregate_claude_runs(runs)
    return aggregate_result, aggregate_parsed, limit_exceeded

def parse_json_file(path: Path) -> tuple[dict[str, Any] | None, str | None]:
    if not path.exists():
        return None, f"missing {path.name}"
    try:
        return json.loads(path.read_text(encoding="utf-8")), None
    except json.JSONDecodeError as exc:
        return None, f"invalid {path.name}: {exc}"


def normalize_string_list(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    return sorted(str(item) for item in value)


def score_blind_profile(payload: dict[str, Any], expected: dict[str, Any]) -> tuple[bool, float, list[str]]:
    score = 0.0
    total = 7.0
    mismatches: list[str] = []
    if payload.get("repoName") == expected.get("repoName"):
        score += 1.0
    else:
        mismatches.append("repoName")
    if payload.get("projectType") == expected.get("projectType"):
        score += 2.0
    else:
        mismatches.append("projectType")
    if payload.get("primaryLanguage") == expected.get("primaryLanguage"):
        score += 1.5
    else:
        mismatches.append("primaryLanguage")
    if sorted(payload.get("packageManagers") or []) == sorted(expected.get("packageManagers") or []):
        score += 1.0
    else:
        mismatches.append("packageManagers")
    payload_dirs = set(payload.get("topLevelDirs") or [])
    expected_dirs = set(expected.get("topLevelDirs") or [])
    overlap = len(payload_dirs & expected_dirs) / max(1, len(expected_dirs))
    if overlap >= 0.5:
        score += 1.0
    else:
        mismatches.append("topLevelDirs")
    if isinstance(payload.get("recommendedRunCommand"), str) and payload["recommendedRunCommand"].strip():
        score += 0.5
    else:
        mismatches.append("recommendedRunCommand")
    blind_score = round(score / total, 3)
    return blind_score >= 0.75, blind_score, mismatches


def load_profile_reports(blind_path: Path, grounded_path: Path, expected: dict[str, Any]) -> dict[str, Any]:
    blind_payload, blind_error = parse_json_file(blind_path)
    grounded_payload, grounded_error = parse_json_file(grounded_path)
    result = {
        "blindProfilePass": False,
        "blindProfileScore": 0.0,
        "blindProfileError": blind_error,
        "blindProfilePayload": blind_payload,
        "groundedProfilePass": False,
        "groundedProfileError": grounded_error,
        "groundedProfilePayload": grounded_payload,
        "profilePass": False,
        "profileError": None,
        "profilePayload": grounded_payload or blind_payload,
    }
    if blind_payload is not None:
        missing = sorted(PROFILE_JSON_KEYS - set(blind_payload))
        if missing:
            result["blindProfileError"] = "missing blind profile keys: " + ", ".join(missing)
        else:
            blind_ok, blind_score, mismatches = score_blind_profile(blind_payload, expected)
            result["blindProfilePass"] = blind_ok
            result["blindProfileScore"] = blind_score
            if mismatches and not blind_ok:
                result["blindProfileError"] = "blind profile mismatches: " + ", ".join(mismatches)
            else:
                result["blindProfileError"] = None
    if grounded_payload is not None:
        missing = sorted(PROFILE_JSON_KEYS - set(grounded_payload))
        if missing:
            result["groundedProfileError"] = "missing grounded profile keys: " + ", ".join(missing)
        else:
            exact_keys = ("repoName", "projectType", "primaryLanguage", "packageManagers", "topLevelDirs", "recommendedRunCommand")
            mismatches = [key for key in exact_keys if grounded_payload.get(key) != expected.get(key)]
            if mismatches:
                result["groundedProfileError"] = "grounded profile mismatches: " + ", ".join(mismatches)
            else:
                result["groundedProfilePass"] = True
                result["groundedProfileError"] = None
    result["profilePass"] = result["blindProfilePass"] and result["groundedProfilePass"]
    if not result["profilePass"]:
        pieces = [value for value in (result["blindProfileError"], result["groundedProfileError"]) if value]
        result["profileError"] = "; ".join(pieces) if pieces else "profile validation failed"
    return result


def load_config_report(path: Path, expected: dict[str, Any]) -> tuple[bool, str | None, dict[str, Any] | None]:
    if not path.exists():
        return False, "missing .ruflo-benchmark/config-report.json", None
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        return False, f"invalid config-report.json: {exc}", None
    missing = sorted(CONFIG_REPORT_KEYS - set(payload))
    if missing:
        return False, f"missing config report keys: {', '.join(missing)}", payload
    payload = dict(payload)
    payload["hookKeys"] = normalize_string_list(payload.get("hookKeys"))
    payload["helperFiles"] = normalize_string_list(payload.get("helperFiles"))
    for key in sorted(CONFIG_REPORT_KEYS):
        if payload.get(key) != expected.get(key):
            return False, f"{key} does not match generated Ruflo config", payload
    return True, None, payload


def extract_hook_commands(settings: dict[str, Any], event_name: str) -> list[str]:
    hooks = settings.get("hooks") or {}
    entries = hooks.get(event_name) or []
    commands: list[str] = []
    for entry in entries:
        for hook in entry.get("hooks") or []:
            command = hook.get("command")
            if isinstance(command, str) and command.strip():
                commands.append(command)
    return commands


def run_verification_commands(workspace: Path, commands: tuple[tuple[str, ...], ...], timeout_seconds: int) -> tuple[bool, list[dict[str, Any]], CommandResult | None]:
    details: list[dict[str, Any]] = []
    final_result: CommandResult | None = None
    for command in commands:
        final_result = run_command(list(command), cwd=workspace, timeout_seconds=timeout_seconds)
        details.append({
            "command": list(command),
            "returncode": final_result.returncode,
            "timedOut": final_result.timed_out,
            "stdout": final_result.stdout,
            "stderr": final_result.stderr,
        })
        if final_result.returncode != 0 or final_result.timed_out:
            return False, details, final_result
    return True, details, final_result


def exercise_runtime_behaviors(root: Path, expected: dict[str, Any]) -> tuple[bool, list[dict[str, Any]], str | None]:
    settings = load_json(root / ".claude" / "settings.json", {}) or {}
    runtime = load_json(root / ".claude-flow" / "config.json", {}) or {}
    checks: list[dict[str, Any]] = []
    nested_cwd = choose_source_dir(root)
    session_id = f"bench-{deterministic_int(root, time.time_ns())}"

    def add_check(name: str, passed: bool, detail: str) -> None:
        checks.append({"name": name, "passed": passed, "detail": detail})

    def run_hook_event(event_name: str, payload: dict[str, Any]) -> list[CommandResult]:
        commands = extract_hook_commands(settings, event_name)
        results: list[CommandResult] = []
        for command in commands:
            result = run_shell_command(command, cwd=nested_cwd, timeout_seconds=20, input_text=json.dumps(payload))
            results.append(result)
        return results

    if expected["hooksEnabled"]:
        pre_results = run_hook_event("PreToolUse", {"hook_event_name": "PreToolUse", "session_id": session_id, "tool_name": "Bash", "tool_input": {"command": "git reset --hard"}})
        deny_pass = any("permissionDecision" in result.stdout and "deny" in result.stdout for result in pre_results)
        add_check("hooks.pretooluse_deny", deny_pass, pre_results[0].stdout if pre_results else "missing PreToolUse command")

        session_start_results = run_hook_event("SessionStart", {"hook_event_name": "SessionStart", "session_id": session_id, "cwd": str(root), "transcript_path": ".claude-flow/sessions/bench-transcript.md"})
        session_current = root / ".claude-flow" / "sessions" / "current.json"
        add_check("hooks.sessionstart", session_current.exists() and all(result.returncode == 0 for result in session_start_results), "current session file created" if session_current.exists() else "current session file missing")

        post_results = run_hook_event("PostToolUse", {"hook_event_name": "PostToolUse", "session_id": session_id, "tool_name": "Write", "tool_input": {"file_path": "tests/benchmark.txt"}})
        events_path = root / ".claude-flow" / "metrics" / "hook-events.jsonl"
        events_text = events_path.read_text(encoding="utf-8") if events_path.exists() else ""
        add_check("hooks.posttooluse", "post-tool-use" in events_text and all(result.returncode == 0 for result in post_results), "post-tool-use event recorded" if "post-tool-use" in events_text else "post-tool-use event missing")

        stop_results = run_hook_event("Stop", {"hook_event_name": "Stop", "session_id": session_id})
        add_check("hooks.stop", all(result.returncode == 0 for result in stop_results), "stop hooks executed" if stop_results else "missing Stop command")

        end_results = run_hook_event("SessionEnd", {"hook_event_name": "SessionEnd", "session_id": session_id, "reason": "benchmark"})
        archived = list((root / ".claude-flow" / "sessions").glob(f"{session_id}.json"))
        add_check("hooks.sessionend", all(result.returncode == 0 for result in end_results) and bool(archived), "session archived" if archived else "session archive missing")

        if expected["agentTeamsEnabled"]:
            task_results = run_hook_event("TaskCompleted", {"hook_event_name": "TaskCompleted", "session_id": session_id})
            idle_results = run_hook_event("TeammateIdle", {"hook_event_name": "TeammateIdle", "session_id": session_id})
            events_text = events_path.read_text(encoding="utf-8") if events_path.exists() else ""
            add_check("hooks.taskcompleted", "task-completed" in events_text and all(result.returncode == 0 for result in task_results), "task-completed event recorded" if "task-completed" in events_text else "task-completed event missing")
            add_check("hooks.teammateidle", "teammate-idle" in events_text and all(result.returncode == 0 for result in idle_results), "teammate-idle event recorded" if "teammate-idle" in events_text else "teammate-idle event missing")
    else:
        add_check("hooks.disabled", "hooks" not in settings or not settings.get("hooks"), "hooks disabled in settings")

    if expected["statuslineEnabled"] and isinstance(settings.get("statusLine"), dict):
        status_command = settings["statusLine"].get("command")
        status_result = run_shell_command(status_command, cwd=nested_cwd, timeout_seconds=20)
        status_ok = f"cf:{expected['profile']}" in status_result.stdout and "hooks:" in status_result.stdout
        add_check("statusline.command", status_ok, status_result.stdout.strip())
    else:
        add_check("statusline.disabled", "statusLine" not in settings, "statusLine omitted from settings")

    if expected["autoMemoryEnabled"]:
        memory_store_rel = ((runtime.get("memory") or {}).get("storePath")) or ".claude-flow/data/auto-memory-store.json"
        memory_store = root / memory_store_rel
        memory_ok = memory_store.exists()
        if memory_ok:
            store = load_json(memory_store, {}) or {}
            meta = store.get("meta") or {}
            memory_ok = bool(meta.get("lastImportAt")) and bool(meta.get("lastSyncAt"))
        add_check("memory.auto_sync", memory_ok, relative_string(root, memory_store) if memory_store.exists() else "memory store missing")
    else:
        add_check("memory.disabled", not (root / ".claude" / "helpers" / "auto-memory-hook.mjs").exists(), "auto-memory helper absent when disabled")

    passed = all(check["passed"] for check in checks)
    failure_detail = None if passed else "; ".join(check["name"] for check in checks if not check["passed"])
    return passed, checks, failure_detail


def exact_pass(row: dict[str, Any]) -> bool:
    return bool(
        row.get("qualityPass")
        and row.get("groundedProfilePass")
        and row.get("configPass")
        and row.get("behaviorPass")
        and row.get("claudeExit") == 0
        and not row.get("claudeTimedOut")
        and not row.get("limitExceeded")
    )


def score_row(row: dict[str, Any]) -> float:
    score = 0.0
    if row["initExit"] == 0:
        score += 15.0
    if row["normalizeExit"] == 0:
        score += 25.0
    if row.get("behaviorPass"):
        score += 45.0
    if row["claudeExit"] == 0:
        score += 70.0
    if row["qualityPass"]:
        score += 95.0
    if row.get("blindProfilePass"):
        score += 20.0
    if row.get("groundedProfilePass"):
        score += 20.0
    if row.get("configPass"):
        score += 35.0
    if row["hookClean"]:
        score += 10.0
    if row["claudeTimedOut"]:
        score -= 150.0
    if isinstance(row.get("claudeDurationMs"), int):
        score -= row["claudeDurationMs"] / 1000.0
    if isinstance(row.get("totalCostUsd"), (int, float)):
        score -= row["totalCostUsd"] * 100.0
    if row.get("cacheBucket") == "warm":
        score -= 1.0
    return round(score, 3)


def diagnose_row(row: dict[str, Any]) -> str:
    if row.get("limitExceeded"):
        return "Claude usage limit hit; waiting for the next reset at 12am America/Anchorage."
    if row.get("claudeTimedOut"):
        return f"Claude timed out after {row.get('claudeElapsedMs')} ms."
    if row.get("initTimedOut"):
        return f"Ruflo init timed out after {row.get('initElapsedMs')} ms."
    if row.get("normalizeTimedOut"):
        return f"Config normalization timed out after {row.get('normalizeElapsedMs')} ms."
    if row.get("failureStage") == "job_budget":
        return "Run exceeded the overall five-minute job budget before all stages completed."
    if row.get("initExit") not in {None, 0}:
        return (row.get("initStderr") or row.get("initStdout") or "Ruflo init failed.").strip().splitlines()[0][:240]
    if row.get("normalizeExit") not in {None, 0}:
        return (row.get("normalizeStderr") or row.get("normalizeStdout") or "Normalization failed.").strip().splitlines()[0][:240]
    if row.get("claudeExit") not in {None, 0}:
        return (row.get("claudeStderr") or row.get("claudeStdout") or "Claude run failed.").strip().splitlines()[0][:240]
    if row.get("verifyExit") not in {None, 0}:
        return (row.get("verifyStderr") or row.get("verifyStdout") or "Verification failed.").strip().splitlines()[0][:240]
    if not row.get("behaviorPass"):
        return row.get("behaviorError") or "runtime behavior validation failed."
    if not row.get("configPass"):
        return row.get("configError") or "config-report.json validation failed."
    if not row.get("profilePass"):
        return row.get("profileError") or "repo profile validation failed."
    return "ok"


def progress_row(row: dict[str, Any]) -> dict[str, Any]:
    return {
        "timestamp": row.get("timestamp"),
        "repo": row.get("repo"),
        "projectType": row.get("projectType"),
        "variant": row.get("variant"),
        "runnerMode": row.get("runnerMode"),
        "scenario": row.get("scenario"),
        "promptStyle": row.get("promptStyle"),
        "cacheBucket": row.get("cacheBucket"),
        "cycle": row.get("cycle"),
        "score": row.get("score"),
        "exactPass": row.get("exactPass"),
        "qualityPass": row.get("qualityPass"),
        "profilePass": row.get("profilePass"),
        "configPass": row.get("configPass"),
        "behaviorPass": row.get("behaviorPass"),
        "hookClean": row.get("hookClean"),
        "claudeTimedOut": row.get("claudeTimedOut"),
        "wallTimeSeconds": row.get("wallTimeSeconds"),
        "failureStage": row.get("failureStage"),
        "diagnosis": row.get("diagnosis"),
    }


def load_existing_results(path: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    if not path.exists():
        return rows
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            rows.append(json.loads(line))
        except json.JSONDecodeError:
            continue
    return rows


def clean_results(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [row for row in rows if not row.get("limitExceeded")]


def confidence_label(runs: int, min_samples: int) -> str:
    if runs >= max(min_samples, 30):
        return "high"
    if runs >= max(6, math.ceil(min_samples * 0.75)):
        return "medium"
    if runs >= 3:
        return "low"
    return "exploratory"


def wilson_lower_bound(successes: int, total: int, z: float = 1.96) -> float:
    if total <= 0:
        return 0.0
    phat = successes / total
    denominator = 1 + z * z / total
    numerator = phat + z * z / (2 * total) - z * math.sqrt((phat * (1 - phat) + z * z / (4 * total)) / total)
    return numerator / denominator


def variant_metrics(rows: list[dict[str, Any]]) -> dict[str, Any]:
    count = len(rows)
    exact_passes = sum(1 for row in rows if row.get("exactPass"))
    quality_passes = sum(1 for row in rows if row["qualityPass"])
    profile_passes = sum(1 for row in rows if row["profilePass"])
    config_passes = sum(1 for row in rows if row.get("configPass"))
    behavior_passes = sum(1 for row in rows if row.get("behaviorPass"))
    hook_clean = sum(1 for row in rows if row["hookClean"])
    timeouts = sum(1 for row in rows if row["claudeTimedOut"])
    cold_rows = [row for row in rows if row.get("cacheBucket") == "cold"]
    warm_rows = [row for row in rows if row.get("cacheBucket") == "warm"]
    return {
        "runs": count,
        "exactPasses": exact_passes,
        "exactRate": round(exact_passes / count, 4) if count else 0.0,
        "wilsonLowerBound": round(wilson_lower_bound(exact_passes, count), 4) if count else 0.0,
        "qualityPasses": quality_passes,
        "qualityRate": round(quality_passes / count, 4) if count else 0.0,
        "profilePasses": profile_passes,
        "profileRate": round(profile_passes / count, 4) if count else 0.0,
        "configPasses": config_passes,
        "configRate": round(config_passes / count, 4) if count else 0.0,
        "behaviorPasses": behavior_passes,
        "behaviorRate": round(behavior_passes / count, 4) if count else 0.0,
        "hookCleanRuns": hook_clean,
        "hookCleanRate": round(hook_clean / count, 4) if count else 0.0,
        "timeouts": timeouts,
        "timeoutRate": round(timeouts / count, 4) if count else 0.0,
        "averageScore": round(sum(row["score"] for row in rows) / count, 3) if count else None,
        "averageCostUsd": round(sum((row["totalCostUsd"] or 0) for row in rows) / count, 6) if count else None,
        "averageDurationMs": round(sum((row["claudeDurationMs"] or 0) for row in rows) / count, 2) if count else None,
        "coldRuns": len(cold_rows),
        "warmRuns": len(warm_rows),
        "coldExactRate": round(sum(1 for row in cold_rows if row.get("exactPass")) / len(cold_rows), 4) if cold_rows else None,
        "warmExactRate": round(sum(1 for row in warm_rows if row.get("exactPass")) / len(warm_rows), 4) if warm_rows else None,
        "runnerModes": sorted({row.get("runnerMode") for row in rows if row.get("runnerMode")}),
        "scenarios": sorted({row.get("scenario") for row in rows if row.get("scenario")}),
    }


def scope_candidate_summary(rows: list[dict[str, Any]], min_samples: int) -> dict[str, dict[str, Any]]:
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in rows:
        grouped[row["variant"]].append(row)
    summary: dict[str, dict[str, Any]] = {}
    for name, records in grouped.items():
        metrics = variant_metrics(records)
        metrics["confidence"] = confidence_label(metrics["runs"], min_samples)
        metrics["meetsSampleFloor"] = metrics["runs"] >= min_samples
        summary[name] = metrics
    return summary


def rank_candidates(candidate_metrics: dict[str, dict[str, Any]], min_samples: int) -> list[str]:
    if not candidate_metrics:
        return []
    any_floor = any(metrics["runs"] >= min_samples for metrics in candidate_metrics.values())

    def key(name: str) -> tuple[Any, ...]:
        metrics = candidate_metrics[name]
        meets_floor = metrics["runs"] >= min_samples if any_floor else True
        return (
            1 if meets_floor else 0,
            metrics["wilsonLowerBound"],
            metrics["exactRate"],
            metrics["behaviorRate"],
            metrics["configRate"],
            -metrics["timeoutRate"],
            metrics["averageScore"] or -10**9,
            -(metrics["averageCostUsd"] or 10**9),
            -(metrics["averageDurationMs"] or 10**9),
            name,
        )

    return sorted(candidate_metrics, key=key, reverse=True)


def recommend_scope(rows: list[dict[str, Any]], min_samples: int, variants_by_name: dict[str, dict[str, Any]]) -> dict[str, Any]:
    candidate_metrics = scope_candidate_summary(rows, min_samples)
    ranked = rank_candidates(candidate_metrics, min_samples)
    winner = ranked[0] if ranked else None
    return {
        "winner": variants_by_name[winner] if winner else None,
        "winnerName": winner,
        "confidence": candidate_metrics[winner]["confidence"] if winner else "exploratory",
        "minSamples": min_samples,
        "candidates": {name: candidate_metrics[name] for name in ranked},
        "rankedVariants": ranked,
    }


def select_variants_for_repo(repo: RepoSpec, results: list[dict[str, Any]], variants: list[Variant], cycle: int, warmup_cycles: int, exploit_variants: int) -> list[Variant]:
    clean = clean_results(results)
    if cycle <= warmup_cycles or not clean:
        return variants

    by_name = {variant.name: variant for variant in variants}
    repo_rows = [row for row in clean if row["repo"] == repo.name]
    type_rows = [row for row in clean if row["projectType"] == repo.project_type]
    repo_counts = Counter(row["variant"] for row in repo_rows)
    chosen: list[str] = []

    under_sampled = sorted((variant.name for variant in variants if repo_counts.get(variant.name, 0) < EXPLORATION_FLOOR), key=lambda name: (repo_counts.get(name, 0), deterministic_int(repo.name, cycle, name)))
    if under_sampled:
        chosen.append(under_sampled[0])

    for rows, min_samples in ((repo_rows, REPO_MIN_SAMPLES), (type_rows, PROJECT_TYPE_MIN_SAMPLES), (clean, GLOBAL_MIN_SAMPLES)):
        metrics = scope_candidate_summary(rows, min_samples)
        for name in rank_candidates(metrics, min_samples):
            if name in by_name and name not in chosen:
                chosen.append(name)
            if len(chosen) >= exploit_variants:
                break
        if len(chosen) >= exploit_variants:
            break

    for variant in sorted(variants, key=lambda item: (repo_counts.get(item.name, 0), deterministic_int(repo.name, cycle, item.name))):
        if variant.name not in chosen:
            chosen.append(variant.name)
        if len(chosen) >= exploit_variants:
            break

    return [by_name[name] for name in chosen[: max(1, exploit_variants)]]


def jobs_for_cycle(repos: list[RepoSpec], results: list[dict[str, Any]], variants: list[Variant], runner_modes: list[str], cycle: int, warmup_cycles: int, exploit_variants: int) -> list[tuple[RepoSpec, Variant, str, str, str]]:
    jobs: list[tuple[RepoSpec, Variant, str, str, str]] = []
    for repo in repos:
        selected_variants = select_variants_for_repo(repo, results, variants, cycle, warmup_cycles, exploit_variants)
        repo_mode_counts = Counter(
            row["runnerMode"]
            for row in results
            if row.get("repo") == repo.name and row.get("runnerMode") in runner_modes
        )
        for variant in selected_variants:
            ordered_modes = sorted(
                runner_modes,
                key=lambda name: (
                    repo_mode_counts.get(name, 0),
                    deterministic_int(repo.name, variant.name, cycle, name),
                ),
            )
            runner_mode = ordered_modes[0]
            repo_mode_counts[runner_mode] += 1
            scenario = choose_scenario(repo, runner_mode, results, cycle)
            prompt_style = choose_prompt_style(repo, scenario.name, runner_mode, results, cycle)
            jobs.append((repo, variant, runner_mode, scenario.name, prompt_style))
    jobs.sort(key=lambda item: deterministic_int(item[0].name, item[1].name, item[2], item[3], item[4], cycle))
    return jobs

def write_progress(*, started: datetime, deadline: datetime, cycle: int, repos: list[RepoSpec], planned_runs: int, completed_runs: int, queued_runs: int, active_jobs: list[dict[str, Any]], recent_rows: list[dict[str, Any]], summary: dict[str, Any] | None, pause_reason: str | None = None, paused_until: str | None = None) -> None:
    payload = {
        "generatedAt": utcnow(),
        "startedAt": started.isoformat(),
        "deadline": deadline.isoformat(),
        "cycle": cycle,
        "repoCount": len(repos),
        "plannedRunsThisCycle": planned_runs,
        "completedRuns": completed_runs,
        "queuedRuns": queued_runs,
        "activeJobs": active_jobs,
        "recentResults": [progress_row(row) for row in recent_rows[-10:]],
        "pauseReason": pause_reason,
        "pausedUntil": paused_until,
        "summary": {
            "totalRuns": summary.get("totalRuns"),
            "cleanRuns": summary.get("cleanRuns"),
            "exactPasses": summary.get("exactPasses"),
            "qualityPasses": summary.get("qualityPasses"),
            "profilePasses": summary.get("profilePasses"),
            "configPasses": summary.get("configPasses"),
            "behaviorPasses": summary.get("behaviorPasses"),
            "hookCleanRuns": summary.get("hookCleanRuns"),
            "timeouts": summary.get("timeouts"),
            "limitExceededRuns": summary.get("limitExceededRuns"),
            "globalDefaultVariant": summary.get("globalDefaultVariant"),
        } if summary else None,
    }
    PROGRESS_PATH.write_text(json.dumps(payload, indent=2) + "\n", encoding="utf-8")


def run_variant(repo: RepoSpec, variant: Variant, runner_mode: str, scenario_name: str, prompt_style: str, temp_root: Path, workspace_mode: str, job_timeout_seconds: int = DEFAULT_JOB_TIMEOUT_SECONDS, previous_results: list[dict[str, Any]] | None = None) -> dict[str, Any]:
    started = time.time()
    started_perf = time.perf_counter()
    workspace = temp_root / f"{repo.name.replace(' ', '_')}__{variant.name}__{runner_mode}__{scenario_name}"
    if workspace.exists():
        shutil.rmtree(workspace, ignore_errors=True)
    workspace.mkdir(parents=True, exist_ok=True)
    workspace_profile = copy_repo(repo, workspace, workspace_mode=workspace_mode)

    init_result = run_ruflo(["init", "--minimal", "--force"], cwd=workspace, timeout_seconds=remaining_timeout(started_perf, job_timeout_seconds, stage_cap=120, reserve_seconds=170))

    normalize_result = CommandResult(command=[], returncode=-2, stdout="", stderr="", timed_out=False, elapsed_ms=0)
    if init_result.returncode == 0 and not init_result.timed_out:
        normalize_result = run_command(
            [
                PYTHON,
                str(RUFLO_NORMALIZER),
                "normalize",
                "--project",
                str(workspace),
                "--profile",
                variant.profile,
                *(["--disable-hooks"] if not variant.hooks else []),
                *(["--disable-statusline"] if not variant.statusline else []),
                *(["--disable-auto-memory"] if not variant.auto_memory else []),
                *(["--agent-teams"] if variant.agent_teams else []),
            ],
            cwd=ROOT,
            timeout_seconds=remaining_timeout(started_perf, job_timeout_seconds, stage_cap=45, reserve_seconds=120),
        )

    repo_context = write_repo_context(workspace, repo, workspace_profile)
    scenario = build_scenario_plan(workspace, repo, scenario_name)
    config_expected = config_expectations(workspace)

    behavior_pass = False
    behavior_checks: list[dict[str, Any]] = []
    behavior_error = None
    if init_result.returncode == 0 and normalize_result.returncode == 0:
        behavior_pass, behavior_checks, behavior_error = exercise_runtime_behaviors(workspace, config_expected)

    claude_result: CommandResult | None = None
    parsed: dict[str, Any] = {}
    verify_result: CommandResult | None = None
    verify_details: list[dict[str, Any]] = []
    profile_report = {
        "blindProfilePass": False,
        "blindProfileScore": 0.0,
        "blindProfileError": None,
        "blindProfilePayload": None,
        "groundedProfilePass": False,
        "groundedProfileError": None,
        "groundedProfilePayload": None,
        "profilePass": False,
        "profileError": None,
        "profilePayload": None,
    }
    config_ok = False
    config_error = None
    config_payload = None
    hook_clean = False
    failure_stage = None
    limit_exceeded = False

    if init_result.timed_out or init_result.returncode != 0:
        failure_stage = "init"
    elif normalize_result.timed_out or normalize_result.returncode != 0:
        failure_stage = "normalize"
    elif (time.perf_counter() - started_perf) >= job_timeout_seconds:
        failure_stage = "job_budget"
    else:
        claude_timeout = remaining_timeout(started_perf, job_timeout_seconds, stage_cap=140, reserve_seconds=20)
        claude_result, parsed, limit_exceeded = run_claude_flow(workspace, repo, scenario, runner_mode, prompt_style, repo_context["recommendedRunCommand"], claude_timeout)
        if claude_result.timed_out:
            failure_stage = "claude"
        elif limit_exceeded:
            failure_stage = "claude"
        elif claude_result.returncode != 0:
            failure_stage = "claude"
        else:
            quality_pass, verify_details, verify_result = run_verification_commands(workspace, scenario.verify_commands, timeout_seconds=remaining_timeout(started_perf, job_timeout_seconds, stage_cap=20))
            if not quality_pass:
                failure_stage = "verify"
        hook_clean = claude_result.returncode == 0 and claude_result.stderr.strip() == ""
        profile_report = load_profile_reports(workspace / ".ruflo-benchmark" / "repo-profile-blind.json", workspace / ".ruflo-benchmark" / "repo-profile-grounded.json", repo_context)
        config_ok, config_error, config_payload = load_config_report(workspace / ".ruflo-benchmark" / "config-report.json", expected=config_expected)
        if not behavior_pass and not failure_stage:
            failure_stage = "behavior"
        elif not config_ok and not failure_stage:
            failure_stage = "config_report"
        elif not profile_report["profilePass"] and not failure_stage:
            failure_stage = "profile"

    prior_context_count = 0
    if previous_results:
        prior_context_count = sum(1 for row in previous_results if row.get("repo") == repo.name and row.get("runnerMode") == runner_mode and row.get("scenario") == scenario_name and row.get("promptStyle") == prompt_style)

    row = {
        "timestamp": utcnow(),
        "repo": repo.name,
        "repoPath": repo.path,
        "projectType": repo.project_type,
        "primaryLanguage": repo.primary_language,
        "fileCount": repo.file_count,
        "variant": variant.name,
        "variantConfig": asdict(variant),
        "runnerMode": runner_mode,
        "scenario": scenario.definition.name,
        "scenarioFamily": scenario.definition.family,
        "scenarioDescription": scenario.definition.description,
        "promptStyle": prompt_style,
        "workspaceMode": workspace_mode,
        "workspaceProfile": workspace_profile,
        "cacheBucket": "cold" if prior_context_count == 0 else "warm",
        "repeatedContextCount": prior_context_count,
        "scenarioFiles": list(scenario.created_files),
        "focusPaths": list(scenario.focus_paths),
        "initExit": init_result.returncode,
        "initTimedOut": init_result.timed_out,
        "initElapsedMs": init_result.elapsed_ms,
        "initStdout": init_result.stdout,
        "initStderr": init_result.stderr,
        "normalizeExit": normalize_result.returncode,
        "normalizeTimedOut": normalize_result.timed_out,
        "normalizeElapsedMs": normalize_result.elapsed_ms,
        "normalizeStdout": normalize_result.stdout,
        "normalizeStderr": normalize_result.stderr,
        "behaviorPass": behavior_pass,
        "behaviorChecks": behavior_checks,
        "behaviorError": behavior_error,
        "claudeExit": claude_result.returncode if claude_result else None,
        "claudeTimedOut": claude_result.timed_out if claude_result else False,
        "claudeElapsedMs": claude_result.elapsed_ms if claude_result else None,
        "claudeStdout": claude_result.stdout if claude_result else "",
        "claudeStderr": claude_result.stderr if claude_result else "",
        "claudeDurationMs": parsed.get("duration_ms"),
        "totalCostUsd": parsed.get("total_cost_usd"),
        "inputTokens": ((parsed.get("usage") or {}).get("input_tokens")),
        "outputTokens": ((parsed.get("usage") or {}).get("output_tokens")),
        "hookClean": hook_clean,
        "limitExceeded": limit_exceeded,
        "qualityPass": verify_result.returncode == 0 if verify_result else False,
        "verifyExit": verify_result.returncode if verify_result else None,
        "verifyTimedOut": verify_result.timed_out if verify_result else False,
        "verifyStdout": verify_result.stdout if verify_result else "",
        "verifyStderr": verify_result.stderr if verify_result else "",
        "verificationDetails": verify_details,
        "configPass": config_ok,
        "configError": config_error,
        "configPayload": config_payload,
        "failureStage": failure_stage,
        "wallTimeSeconds": round(time.time() - started, 3),
    }
    row.update(profile_report)
    row["exactPass"] = exact_pass(row)
    row["diagnosis"] = diagnose_row(row)
    row["score"] = score_row(row)
    shutil.rmtree(workspace, ignore_errors=True)
    return row


def build_recommendations(results: list[dict[str, Any]], runner_modes: list[str]) -> dict[str, Any]:
    variants_by_name = {variant.name: asdict(variant) for variant in VARIANTS}
    claude_rows = [row for row in clean_results(results) if row["claudeExit"] is not None]
    overall = recommend_scope(claude_rows, GLOBAL_MIN_SAMPLES, variants_by_name)

    project_type_defaults: dict[str, dict[str, Any]] = {}
    project_type_evidence: dict[str, Any] = {}
    for project_type in sorted({row["projectType"] for row in claude_rows}):
        rows = [row for row in claude_rows if row["projectType"] == project_type]
        evidence = recommend_scope(rows, PROJECT_TYPE_MIN_SAMPLES, variants_by_name)
        project_type_evidence[project_type] = evidence
        if evidence["winner"]:
            project_type_defaults[project_type] = evidence["winner"]

    repo_overrides: dict[str, dict[str, Any]] = {}
    repo_override_evidence: dict[str, Any] = {}
    for repo_name in sorted({row["repo"] for row in claude_rows}):
        rows = [row for row in claude_rows if row["repo"] == repo_name]
        evidence = recommend_scope(rows, REPO_MIN_SAMPLES, variants_by_name)
        repo_override_evidence[repo_name] = evidence
        ptype = next((row["projectType"] for row in rows), None)
        type_default = project_type_defaults.get(ptype, {}).get("name")
        if evidence["winner"] and evidence["winnerName"] and evidence["winnerName"] != type_default:
            repo_overrides[repo_name] = evidence["winner"]

    by_runner_mode: dict[str, Any] = {}
    for runner_mode in runner_modes:
        mode_rows = [row for row in claude_rows if row.get("runnerMode") == runner_mode]
        mode_recommendation = recommend_scope(mode_rows, GLOBAL_MIN_SAMPLES, variants_by_name)
        by_runner_mode[runner_mode] = {"globalDefault": mode_recommendation["winner"], "globalEvidence": mode_recommendation}

    return {
        "generatedAt": utcnow(),
        "rankingMethod": {
            "winnerSelection": "wilson_lower_bound_over_exact_pass_rate_then_score_cost_duration",
            "exactPassFields": ["qualityPass", "profilePass", "configPass", "behaviorPass", "claudeExit==0", "noTimeout", "noLimitExceeded"],
            "sampleFloors": {"global": GLOBAL_MIN_SAMPLES, "projectType": PROJECT_TYPE_MIN_SAMPLES, "repo": REPO_MIN_SAMPLES},
            "runnerModes": runner_modes,
            "workspaceMode": DEFAULT_WORKSPACE_MODE,
        },
        "globalDefault": overall["winner"],
        "globalEvidence": overall,
        "projectTypeDefaults": project_type_defaults,
        "projectTypeEvidence": project_type_evidence,
        "repoOverrides": repo_overrides,
        "repoOverrideEvidence": repo_override_evidence,
        "byRunnerMode": by_runner_mode,
    }


def build_summary(results: list[dict[str, Any]], repos: list[RepoSpec], cycle: int, runner_modes: list[str]) -> dict[str, Any]:
    clean = clean_results(results)
    claude_rows = [row for row in clean if row["claudeExit"] is not None]
    recommendations = build_recommendations(results, runner_modes)
    global_variant = recommendations.get("globalDefault", {}) or {}
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in claude_rows:
        grouped[row["variant"]].append(row)
    variant_board = {name: variant_metrics(grouped[name]) for name in sorted(grouped)}
    scenario_grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in claude_rows:
        scenario_grouped[row["scenario"]].append(row)
    scenario_board = {name: variant_metrics(scenario_grouped[name]) for name in sorted(scenario_grouped)}
    return {
        "generatedAt": utcnow(),
        "cycle": cycle,
        "repos": [asdict(repo) for repo in repos],
        "runnerModes": runner_modes,
        "totalRuns": len(results),
        "cleanRuns": len(clean),
        "successfulClaudeRuns": sum(1 for row in claude_rows if row["claudeExit"] == 0),
        "exactPasses": sum(1 for row in claude_rows if row.get("exactPass")),
        "qualityPasses": sum(1 for row in claude_rows if row["qualityPass"]),
        "profilePasses": sum(1 for row in claude_rows if row["profilePass"]),
        "blindProfilePasses": sum(1 for row in claude_rows if row.get("blindProfilePass")),
        "groundedProfilePasses": sum(1 for row in claude_rows if row.get("groundedProfilePass")),
        "configPasses": sum(1 for row in claude_rows if row.get("configPass")),
        "behaviorPasses": sum(1 for row in claude_rows if row.get("behaviorPass")),
        "hookCleanRuns": sum(1 for row in claude_rows if row["hookClean"]),
        "timeouts": sum(1 for row in claude_rows if row["claudeTimedOut"]),
        "globalDefaultVariant": global_variant.get("name"),
        "globalDefault": global_variant,
        "globalEvidence": recommendations.get("globalEvidence"),
        "variantLeaderboard": variant_board,
        "scenarioLeaderboard": scenario_board,
        "projectTypeDefaults": recommendations.get("projectTypeDefaults", {}),
        "projectTypeEvidence": recommendations.get("projectTypeEvidence", {}),
        "repoOverrides": recommendations.get("repoOverrides", {}),
        "repoOverrideEvidence": recommendations.get("repoOverrideEvidence", {}),
        "byRunnerMode": recommendations.get("byRunnerMode", {}),
        "recentFailures": [progress_row(row) for row in results if row.get("failureStage")][-10:],
        "limitExceededRuns": sum(1 for row in results if row.get("limitExceeded")),
    }

def write_docx(summary: dict[str, Any], results: list[dict[str, Any]], started_at: str, deadline: str) -> None:
    doc = Document()
    doc.add_heading("Ruflo Repo Optimizer", 0)
    doc.add_paragraph(f"Started: {started_at}")
    doc.add_paragraph(f"Deadline: {deadline}")
    doc.add_paragraph("Method: clone real repos into isolated workspaces, initialize Ruflo, normalize a candidate runtime variant, exercise hook/statusline/memory behavior directly, run benchmark scenarios inside the repo tree, score blind and grounded repo understanding separately, and aggregate normalized recommendation evidence with sample floors.")

    doc.add_heading("Current Summary", level=1)
    for key in ("cycle", "totalRuns", "cleanRuns", "successfulClaudeRuns", "exactPasses", "qualityPasses", "profilePasses", "configPasses", "behaviorPasses", "hookCleanRuns", "timeouts", "limitExceededRuns", "globalDefaultVariant"):
        doc.add_paragraph(f"{key}: {summary.get(key)}")

    doc.add_heading("Global Evidence", level=1)
    doc.add_paragraph(json.dumps(summary.get("globalEvidence") or {}, ensure_ascii=False))

    doc.add_heading("Project Type Defaults", level=1)
    for name, payload in (summary.get("projectTypeDefaults") or {}).items():
        doc.add_paragraph(f"{name}: {payload}")

    doc.add_heading("Repo Overrides", level=1)
    for name, payload in (summary.get("repoOverrides") or {}).items():
        doc.add_paragraph(f"{name}: {payload}")

    doc.add_heading("Variant Leaderboard", level=1)
    for name, payload in (summary.get("variantLeaderboard") or {}).items():
        doc.add_paragraph(f"{name}: {payload}")

    doc.add_heading("Scenario Leaderboard", level=1)
    for name, payload in (summary.get("scenarioLeaderboard") or {}).items():
        doc.add_paragraph(f"{name}: {payload}")

    doc.add_heading("Recent Failures", level=1)
    for row in summary.get("recentFailures") or []:
        doc.add_paragraph(json.dumps(row, ensure_ascii=False))

    doc.add_heading("Recent Runs", level=1)
    for row in results[-12:]:
        doc.add_paragraph(json.dumps({
            "repo": row["repo"],
            "variant": row["variant"],
            "runnerMode": row.get("runnerMode"),
            "scenario": row.get("scenario"),
            "cacheBucket": row.get("cacheBucket"),
            "score": row["score"],
            "exactPass": row.get("exactPass"),
            "qualityPass": row["qualityPass"],
            "profilePass": row["profilePass"],
            "configPass": row.get("configPass"),
            "behaviorPass": row.get("behaviorPass"),
            "claudeDurationMs": row["claudeDurationMs"],
            "totalCostUsd": row["totalCostUsd"],
        }, ensure_ascii=False))

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Adaptive Ruflo optimizer for real local repos.")
    parser.add_argument("--projects-root", type=Path, default=PROJECTS_ROOT)
    parser.add_argument("--repos", nargs="*", default=None)
    parser.add_argument("--max-repos", type=int, default=10)
    parser.add_argument("--variant-limit", type=int, default=len(VARIANTS))
    parser.add_argument("--duration-hours", type=float, default=12.0)
    parser.add_argument("--max-cycles", type=int, default=0)
    parser.add_argument("--results-root", type=Path, default=RESULTS_ROOT)
    parser.add_argument("--max-parallel", type=int, default=DEFAULT_MAX_PARALLEL)
    parser.add_argument("--job-timeout-seconds", type=int, default=DEFAULT_JOB_TIMEOUT_SECONDS)
    parser.add_argument("--warmup-cycles", type=int, default=DEFAULT_WARMUP_CYCLES)
    parser.add_argument("--exploit-variants", type=int, default=DEFAULT_EXPLOIT_VARIANTS)
    parser.add_argument("--runner-modes", nargs="*", default=list(DEFAULT_RUNNER_MODES), choices=sorted(RUNNER_MODES))
    parser.add_argument("--workspace-mode", choices=["realistic", "compact"], default=DEFAULT_WORKSPACE_MODE)
    parser.add_argument("--fresh", action="store_true")
    return parser.parse_args()


def main() -> int:
    global RESULTS_ROOT, PROGRESS_PATH, SUMMARY_PATH, RESULTS_PATH, RECOMMENDATIONS_PATH, DOCX_PATH, WORKER_STATE_PATH
    args = parse_args()
    RESULTS_ROOT = args.results_root.resolve()
    PROGRESS_PATH = RESULTS_ROOT / "progress.json"
    SUMMARY_PATH = RESULTS_ROOT / "summary.json"
    RESULTS_PATH = RESULTS_ROOT / "results.jsonl"
    RECOMMENDATIONS_PATH = RESULTS_ROOT / "recommended-settings.json"
    DOCX_PATH = RESULTS_ROOT / "ruflo-repo-optimizer.docx"
    WORKER_STATE_PATH = RESULTS_ROOT / "worker.json"
    repos = discover_repos(args.projects_root, args.repos, args.max_repos)
    if not repos:
        raise SystemExit("No project repos found to optimize.")
    variants = VARIANTS[: max(1, min(args.variant_limit, len(VARIANTS)))]
    runner_modes = [mode for mode in args.runner_modes if mode in RUNNER_MODES]
    if not runner_modes:
        raise SystemExit("No runner modes selected.")

    if args.fresh and RESULTS_ROOT.exists():
        shutil.rmtree(RESULTS_ROOT, ignore_errors=True)
    RESULTS_ROOT.mkdir(parents=True, exist_ok=True)
    existing_progress = {}
    if not args.fresh and PROGRESS_PATH.exists():
        try:
            existing_progress = json.loads(PROGRESS_PATH.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            existing_progress = {}
    existing_summary = {}
    if not args.fresh and SUMMARY_PATH.exists():
        try:
            existing_summary = json.loads(SUMMARY_PATH.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            existing_summary = {}
    started = datetime.fromisoformat(existing_progress["startedAt"]) if existing_progress.get("startedAt") else datetime.now(timezone.utc)
    deadline = datetime.fromisoformat(existing_progress["deadline"]) if existing_progress.get("deadline") else started + timedelta(hours=args.duration_hours)
    cycle = int(existing_summary.get("cycle", 0) or 0)
    results: list[dict[str, Any]] = load_existing_results(RESULTS_PATH) if not args.fresh else []

    WORKER_STATE_PATH.write_text(json.dumps({
        "pid": os.getpid(),
        "startedAt": utcnow(),
        "resultsRoot": str(RESULTS_ROOT),
        "argv": sys.argv,
        "runnerModes": runner_modes,
        "workspaceMode": args.workspace_mode,
    }, indent=2) + "\n", encoding="utf-8")

    if datetime.now(timezone.utc) >= deadline:
        final_summary = build_summary(results, repos, cycle, runner_modes)
        SUMMARY_PATH.write_text(json.dumps(final_summary, indent=2) + "\n", encoding="utf-8")
        RECOMMENDATIONS_PATH.write_text(json.dumps(build_recommendations(results, runner_modes), indent=2) + "\n", encoding="utf-8")
        write_docx(final_summary, results, started.isoformat(), deadline.isoformat())
        return 0

    with tempfile.TemporaryDirectory(prefix="ruflo-repo-optimizer-") as temp_root_raw:
        temp_root = Path(temp_root_raw)
        while datetime.now(timezone.utc) < deadline:
            cycle += 1
            cycle_jobs = jobs_for_cycle(repos, results, variants, runner_modes, cycle, warmup_cycles=max(0, args.warmup_cycles), exploit_variants=max(1, min(args.exploit_variants, len(variants))))
            pending_jobs = list(cycle_jobs)
            futures: dict[Any, tuple[RepoSpec, Variant, str, str, str]] = {}
            active_jobs: dict[Any, dict[str, Any]] = {}
            summary = build_summary(results, repos, cycle - 1, runner_modes) if results else None
            limit_pause_until: datetime | None = None
            write_progress(started=started, deadline=deadline, cycle=cycle, repos=repos, planned_runs=len(cycle_jobs), completed_runs=len(results), queued_runs=len(pending_jobs), active_jobs=[], recent_rows=results, summary=summary, pause_reason=None, paused_until=None)

            with ThreadPoolExecutor(max_workers=max(1, args.max_parallel)) as executor:
                while pending_jobs or futures:
                    if limit_pause_until and datetime.now(timezone.utc) < limit_pause_until and not futures:
                        summary = build_summary(results, repos, cycle - 1 if not results else cycle, runner_modes)
                        write_progress(started=started, deadline=deadline, cycle=cycle, repos=repos, planned_runs=len(cycle_jobs), completed_runs=len(results), queued_runs=len(pending_jobs), active_jobs=[], recent_rows=results, summary=summary, pause_reason="claude_limit_reached", paused_until=limit_pause_until.isoformat())
                        time.sleep(min(60, max(5, int((limit_pause_until - datetime.now(timezone.utc)).total_seconds()))))
                        continue

                    while pending_jobs and len(futures) < max(1, args.max_parallel):
                        if limit_pause_until and datetime.now(timezone.utc) < limit_pause_until:
                            break
                        repo, variant, runner_mode, scenario_name, prompt_style = pending_jobs.pop(0)
                        future = executor.submit(run_variant, repo, variant, runner_mode, scenario_name, prompt_style, temp_root, args.workspace_mode, max(30, args.job_timeout_seconds), list(results))
                        futures[future] = (repo, variant, runner_mode, scenario_name, prompt_style)
                        active_jobs[future] = {"repo": repo.name, "projectType": repo.project_type, "variant": variant.name, "runnerMode": runner_mode, "scenario": scenario_name, "promptStyle": prompt_style, "submittedAt": utcnow()}

                    summary = build_summary(results, repos, cycle - 1 if not results else cycle, runner_modes)
                    write_progress(started=started, deadline=deadline, cycle=cycle, repos=repos, planned_runs=len(cycle_jobs), completed_runs=len(results), queued_runs=len(pending_jobs), active_jobs=list(active_jobs.values()), recent_rows=results, summary=summary, pause_reason="claude_limit_reached" if limit_pause_until and datetime.now(timezone.utc) < limit_pause_until else None, paused_until=limit_pause_until.isoformat() if limit_pause_until and datetime.now(timezone.utc) < limit_pause_until else None)

                    if not futures:
                        continue

                    done, _ = wait(set(futures), timeout=2.0, return_when=FIRST_COMPLETED)
                    if not done:
                        continue

                    for future in done:
                        repo, variant, runner_mode, scenario_name, prompt_style = futures.pop(future)
                        active_jobs.pop(future, None)
                        try:
                            row = future.result()
                        except Exception as exc:
                            row = {
                                "timestamp": utcnow(),
                                "repo": repo.name,
                                "repoPath": repo.path,
                                "projectType": repo.project_type,
                                "primaryLanguage": repo.primary_language,
                                "fileCount": repo.file_count,
                                "variant": variant.name,
                                "variantConfig": asdict(variant),
                                "runnerMode": runner_mode,
                                "scenario": scenario_name,
                                "promptStyle": prompt_style,
                                "workspaceMode": args.workspace_mode,
                                "cacheBucket": "cold",
                                "repeatedContextCount": 0,
                                "initExit": None,
                                "initTimedOut": False,
                                "initElapsedMs": None,
                                "initStdout": "",
                                "initStderr": "",
                                "normalizeExit": None,
                                "normalizeTimedOut": False,
                                "normalizeElapsedMs": None,
                                "normalizeStdout": "",
                                "normalizeStderr": "",
                                "behaviorPass": False,
                                "behaviorChecks": [],
                                "behaviorError": str(exc),
                                "claudeExit": None,
                                "claudeTimedOut": False,
                                "claudeElapsedMs": None,
                                "claudeStdout": "",
                                "claudeStderr": "",
                                "claudeDurationMs": None,
                                "totalCostUsd": None,
                                "inputTokens": None,
                                "outputTokens": None,
                                "hookClean": False,
                                "limitExceeded": False,
                                "qualityPass": False,
                                "verifyExit": None,
                                "verifyTimedOut": False,
                                "verifyStdout": "",
                                "verifyStderr": "",
                                "configPass": False,
                                "configError": str(exc),
                                "configPayload": None,
                                "blindProfilePass": False,
                                "blindProfileScore": 0.0,
                                "blindProfileError": str(exc),
                                "blindProfilePayload": None,
                                "groundedProfilePass": False,
                                "groundedProfileError": str(exc),
                                "groundedProfilePayload": None,
                                "profilePass": False,
                                "profileError": str(exc),
                                "profilePayload": None,
                                "failureStage": "scheduler",
                                "wallTimeSeconds": 0.0,
                                "exactPass": False,
                            }
                            row["diagnosis"] = f"Worker raised {type(exc).__name__}: {exc}"
                            row["score"] = -500.0
                        row["cycle"] = cycle
                        results.append(row)
                        if row.get("limitExceeded"):
                            limit_pause_until = next_anchorage_reset()

                        with RESULTS_PATH.open("a", encoding="utf-8") as handle:
                            handle.write(json.dumps(row) + "\n")

                        recommendations = build_recommendations(results, runner_modes)
                        summary = build_summary(results, repos, cycle, runner_modes)
                        SUMMARY_PATH.write_text(json.dumps(summary, indent=2) + "\n", encoding="utf-8")
                        RECOMMENDATIONS_PATH.write_text(json.dumps(recommendations, indent=2) + "\n", encoding="utf-8")
                        write_docx(summary, results, started.isoformat(), deadline.isoformat())
                        write_progress(started=started, deadline=deadline, cycle=cycle, repos=repos, planned_runs=len(cycle_jobs), completed_runs=len(results), queued_runs=len(pending_jobs), active_jobs=list(active_jobs.values()), recent_rows=results, summary=summary, pause_reason="claude_limit_reached" if limit_pause_until and datetime.now(timezone.utc) < limit_pause_until else None, paused_until=limit_pause_until.isoformat() if limit_pause_until and datetime.now(timezone.utc) < limit_pause_until else None)

            if args.max_cycles and cycle >= args.max_cycles:
                break

    final_summary = build_summary(results, repos, cycle, runner_modes)
    SUMMARY_PATH.write_text(json.dumps(final_summary, indent=2) + "\n", encoding="utf-8")
    RECOMMENDATIONS_PATH.write_text(json.dumps(build_recommendations(results, runner_modes), indent=2) + "\n", encoding="utf-8")
    write_docx(final_summary, results, started.isoformat(), deadline.isoformat())
    WORKER_STATE_PATH.write_text(json.dumps({"pid": os.getpid(), "startedAt": utcnow(), "resultsRoot": str(RESULTS_ROOT), "argv": sys.argv, "runnerModes": runner_modes, "workspaceMode": args.workspace_mode}, indent=2) + "\n", encoding="utf-8")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
